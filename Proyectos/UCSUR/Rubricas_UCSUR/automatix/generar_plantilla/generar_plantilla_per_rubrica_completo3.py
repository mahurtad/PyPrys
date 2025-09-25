import os
import re
import pandas as pd
from docx import Document
from openpyxl import load_workbook

# === CONFIGURACIÓN ===
base_folder_path = r"C:\Users\manue\Downloads\Medicina Veterinaria y Zootecnia - Copy\Tecnología e Industrialización de Alimentos\New folder"
error_log = []
validation_log = []
document_log = []

# === MAPA DE EQUIVALENCIAS (ampliado) ===
mapa_equivalencias = {
    "EC1": "Evaluación Continua 1", "EC2": "Evaluación Continua 2", 
    "EC3": "Evaluación Continua 3", "EC4": "Evaluación Continua 4",
    "EF": "Evaluación Final", "EP": "Evaluación Parcial", 
    "ED": "Evaluación Diagnóstica",
    "CONTINUA1": "Evaluación Continua 1", "CONTINUA2": "Evaluación Continua 2", 
    "CONTINUA3": "Evaluación Continua 3", "CONTINUA4": "Evaluación Continua 4", 
    "DIAGNOSTICA": "Evaluación Diagnóstica", 
    "FINAL": "Evaluación Final", "PARCIAL": "Evaluación Parcial",
    "EVFINAL": "Evaluación Final (ES-FINAL)", "EV DIAG": "Evaluación Diagnóstica (Ficha)"
}

# === NUEVAS FUNCIONES DE DETECCIÓN DE FORMATOS ===
def detectar_formato_rubrica(doc):
    """Identifica el formato de la rúbrica basado en patrones específicos"""
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 4:
                texto = row.cells[0].text.strip().upper()
                if "CRITERIO" in texto or "ASPECTO" in texto:
                    # Verificar si es formato EC1 (puntos entre paréntesis)
                    if any("(" in cell.text and "puntos" in cell.text.lower() for cell in row.cells[1:]):
                        return "FORMATO_EC1"
                    # Verificar si es formato EC2/EC3 (puntos con asteriscos)
                    elif any("puntos)" in cell.text.lower() and "*" in cell.text for cell in row.cells[1:]):
                        return "FORMATO_EC2"
                    # Verificar si es formato de dos filas
                    elif len(table.rows) >= 2 and "PUNTOS" not in row.cells[1].text.upper():
                        siguiente_fila = table.rows[table.rows.index(row)+1]
                        if any("puntos" in cell.text.lower() for cell in siguiente_fila.cells[1:]):
                            return "FORMATO_DOS_FILAS"
    return "FORMATO_DESCONOCIDO"

def normalizar_puntaje(valor):
    """Versión mejorada que maneja más formatos de puntaje"""
    if isinstance(valor, str):
        valor = valor.strip().replace(",", ".")
        # Manejar casos como "5 puntos)", "(5 pts", "20*", etc.
        valor = re.sub(r"[^\d.]", "", valor)
    try:
        return float(valor) if valor else None
    except (ValueError, TypeError):
        return None

def extraer_puntaje_y_descripcion_mejorado(texto):
    """Versión mejorada que maneja los formatos de EC1, EC2 y EC3"""
    texto = texto.strip()
    
    # Caso EC1: "(5 puntos)" al final
    match_ec1 = re.search(r"(.*?)\s*\((\d+[.,]?\d*)\s*puntos\)\s*$", texto, re.IGNORECASE)
    if match_ec1:
        return normalizar_puntaje(match_ec1.group(2)), match_ec1.group(1).strip()
    
    # Caso EC2/EC3: "(20 puntos)***" o similar
    match_ec2 = re.search(r"\((\d+[.,]?\d*)\s*puntos\)\*+\s*$", texto, re.IGNORECASE)
    if match_ec2:
        return normalizar_puntaje(match_ec2.group(1)), texto.replace(match_ec2.group(0), "").strip()
    
    # Caso genérico: número seguido de "pts" o "puntos"
    match_generico = re.search(r"(\d+[.,]?\d*)\s*(?:pts|puntos)\b", texto, re.IGNORECASE)
    if match_generico:
        return normalizar_puntaje(match_generico.group(1)), texto.replace(match_generico.group(0), "").strip()
    
    # Si no encuentra patrón, devolver todo como descripción
    return None, texto

# === PROCESAMIENTO DE DOCUMENTO ACTUALIZADO ===
def procesar_documento(word_file_path):
    try:
        doc = Document(word_file_path)
        evaluacion_data = []
        formato = detectar_formato_rubrica(doc)
        
        for table in doc.tables:
            if formato == "FORMATO_DOS_FILAS":
                # Procesamiento para tablas con descripción y puntaje en filas separadas
                for i in range(0, len(table.rows) - 1, 2):
                    if i + 1 >= len(table.rows):
                        continue
                        
                    desc_row = table.rows[i]
                    puntaje_row = table.rows[i + 1]
                    
                    if len(desc_row.cells) < 4 or len(puntaje_row.cells) < 4:
                        continue
                        
                    criterio = desc_row.cells[0].text.strip()
                    if es_fila_cabecera(criterio) or es_fila_resumen(criterio):
                        continue
                        
                    desc_ld = desc_row.cells[1].text.strip()
                    desc_l = desc_row.cells[2].text.strip()
                    desc_nl = desc_row.cells[3].text.strip()
                    
                    p_ld = extraer_puntaje_desde_texto(puntaje_row.cells[1].text.strip())
                    p_l = extraer_puntaje_desde_texto(puntaje_row.cells[2].text.strip())
                    p_nl = extraer_puntaje_desde_texto(puntaje_row.cells[3].text.strip())
                    
                    if p_ld is not None and p_l is not None and p_nl is not None:
                        evaluacion_data.append([
                            criterio,
                            p_ld, "Logro Destacado", desc_ld,
                            p_l, "Logrado", desc_l,
                            p_nl, "No Logrado", desc_nl
                        ])
            
            else:  # FORMATO_EC1, FORMATO_EC2 o desconocido
                for row in table.rows:
                    if len(row.cells) < 4:
                        continue
                        
                    celda_0 = row.cells[0].text.strip()
                    if es_fila_cabecera(celda_0) or es_fila_resumen(celda_0) or not celda_0:
                        continue
                        
                    celda_1 = row.cells[1].text.strip()
                    celda_2 = row.cells[2].text.strip()
                    celda_3 = row.cells[3].text.strip()
                    
                    p_ld, desc_ld = extraer_puntaje_y_descripcion_mejorado(celda_1)
                    p_l, desc_l = extraer_puntaje_y_descripcion_mejorado(celda_2)
                    p_nl, desc_nl = extraer_puntaje_y_descripcion_mejorado(celda_3)
                    
                    # Si no se extrajo puntaje pero hay texto, usar el texto como descripción
                    desc_ld = desc_ld if p_ld is not None else celda_1
                    desc_l = desc_l if p_l is not None else celda_2
                    desc_nl = desc_nl if p_nl is not None else celda_3
                    
                    # Si no se extrajo puntaje pero el texto contiene "logro destacado", etc.
                    if p_ld is None and "destacado" in celda_1.lower():
                        p_ld = 5.0  # Valor por defecto para EC1
                    if p_l is None and "logrado" in celda_2.lower():
                        p_l = 3.75
                    if p_nl is None and "no logrado" in celda_3.lower():
                        p_nl = 3.0
                    
                    if p_ld is not None and p_l is not None and p_nl is not None:
                        evaluacion_data.append([
                            celda_0,
                            p_ld, "Logro Destacado", desc_ld,
                            p_l, "Logrado", desc_l,
                            p_nl, "No Logrado", desc_nl
                        ])

        if not evaluacion_data:
            error_log.append([word_file_path, "No se encontraron datos válidos en el documento"])
            return None
            
        df = pd.DataFrame(evaluacion_data, columns=[
            "ASPECTO / CRITERIO A EVALUAR", 
            "Puntaje de la calificación", "Título de la calificación", "Descripción de la calificación .1", 
            "Puntaje de la calificación.1", "Título de la calificación .1", "Descripción de la calificación .1.1", 
            "Puntaje de la calificación.2", "Título de la calificación .2", "Descripción de la calificación .2"
        ])
        
        # Validación flexible de suma total (diferentes evaluaciones tienen distintos totales)
        suma_ld = df["Puntaje de la calificación"].sum()
        if not (15 <= suma_ld <= 25):  # Rango amplio para diferentes evaluaciones
            error_log.append([word_file_path, f"La suma de 'Logro Destacado' es inusual: {suma_ld:.2f}"])
        
        return df

    except Exception as e:
        error_log.append([word_file_path, f"Error al procesar el documento: {str(e)}"])
        return None

# === FUNCIONES AUXILIARES (sin cambios) ===
def es_fila_cabecera(texto):
    return bool(re.search(r"ASPECTO|CRITERIO|EVALUAR|COMPETENCIAS|LOGRO DESTACADO|LOGRADO|NO LOGRADO", texto.upper()))

def es_fila_resumen(texto):
    return bool(re.search(r"\d+\s*[-–]\s*\d+\s*(?:pts|puntos)|TOTAL|PUNTAJE", texto.upper()))

def extraer_puntaje_desde_texto(texto):
    texto = texto.strip()
    match = re.search(r"(\d+[.,]?\d*)", texto)
    if match:
        return normalizar_puntaje(match.group(1))
    return None

def detectar_tipo_evaluacion(nombre_archivo):
    nombre_upper = nombre_archivo.upper()
    for clave in mapa_equivalencias.keys():
        if clave in nombre_upper:
            return mapa_equivalencias[clave]
    return None

# === PROCESAMIENTO PRINCIPAL (sin cambios) ===
def procesar_carpetas():
    for root, dirs, files in os.walk(base_folder_path):
        word_files = [f for f in files if f.endswith(".docx") and not f.startswith("~$")]
        if not word_files:
            continue
            
        folder_name = os.path.basename(root)
        output_excel_path = os.path.join(root, f"{folder_name}.xlsx")
        dfs_por_hoja = {}
        
        for filename in word_files:
            word_file_path = os.path.join(root, filename)
            tipo_eval = detectar_tipo_evaluacion(filename)
            
            if not tipo_eval:
                error_log.append([word_file_path, "Tipo de evaluación no reconocido"])
                document_log.append([filename, "Ignorado (tipo no reconocido)"])
                continue
                
            df = procesar_documento(word_file_path)
            if df is not None:
                sheet_name = tipo_eval[:31]
                
                if sheet_name in dfs_por_hoja:
                    counter = 1
                    while f"{sheet_name}_{counter}" in dfs_por_hoja:
                        counter += 1
                    sheet_name = f"{sheet_name}_{counter}"
                
                dfs_por_hoja[sheet_name] = df
                document_log.append([filename, f"Procesado correctamente -> {sheet_name}"])
            else:
                document_log.append([filename, "Procesado con errores"])
        
        if dfs_por_hoja:
            try:
                with pd.ExcelWriter(output_excel_path, engine='openpyxl') as writer:
                    for sheet_name, df in dfs_por_hoja.items():
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                print(f"✅ Archivo Excel generado: {output_excel_path}")
                
                validation_log.append([
                    output_excel_path,
                    ", ".join(word_files),
                    ", ".join(dfs_por_hoja.keys()),
                    "OK" if len(dfs_por_hoja) == len([f for f in document_log if "correctamente" in f[1]]) else "Algunos archivos no se procesaron"
                ])
            except Exception as e:
                error_log.append([output_excel_path, f"Error al generar Excel: {str(e)}"])

# === EJECUCIÓN Y REPORTES (sin cambios) ===
if __name__ == "__main__":
    print("Iniciando procesamiento de documentos...")
    procesar_carpetas()
    
    # Generar reportes
    if error_log:
        pd.DataFrame(error_log, columns=["Archivo", "Error"]).to_excel(
            os.path.join(base_folder_path, "Errores_de_procesamiento.xlsx"), index=False)
        
    if validation_log:
        pd.DataFrame(validation_log, columns=["Archivo Excel", "Archivos Word", "Hojas en Excel", "Observación"]).to_excel(
            os.path.join(base_folder_path, "Resumen_de_validacion.xlsx"), index=False)
        
    if document_log:
        pd.DataFrame(document_log, columns=["Archivo", "Estado"]).to_excel(
            os.path.join(base_folder_path, "Resumen_de_procesamiento.xlsx"), index=False)
    
    print("\nResumen:")
    print(f"- Documentos procesados correctamente: {len([x for x in document_log if 'correctamente' in x[1]])}")
    print(f"- Documentos con errores: {len([x for x in document_log if 'errores' in x[1]])}")
    print(f"- Errores registrados: {len(error_log)}")
    print("\nProceso completado. Revise los archivos de reporte si hay errores.")