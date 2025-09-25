import os
from docx import Document

# Lista de cursos
cursos = [
    "Biologia",
    "Quimica general",
    "Química Orgánica",
    "Matemática",
    "Matematica general",
    "Matematica I",
    "Matematica II",
    "Álgebra",
    "Matematica III",
    "Bioquimica",
    "Estadistica General",
    "Fisica",
    "Fisica I"
]

# Subcarpetas relevantes para insertar documento Word
subcarpetas_target = {
    "Videos del curso": "Enlaces de Videos.docx",
    "Materiales adicionales": "Enlaces de Materiales adicionales.docx"
}

# Ruta base
ruta_base = r"C:\Users\manue\OneDrive - Grupo Educad\Eventos de Innovación Educativa Docente - Materiales - tutores virtuales CCBB"

# Crear documentos en las subcarpetas correspondientes
for curso in cursos:
    for subcarpeta, nombre_doc in subcarpetas_target.items():
        ruta_sub = os.path.join(ruta_base, curso, subcarpeta)
        if os.path.exists(ruta_sub):
            doc_path = os.path.join(ruta_sub, nombre_doc)
            if not os.path.exists(doc_path):
                doc = Document()
                doc.add_heading(nombre_doc.replace(".docx", ""), level=1)
                doc.add_paragraph("Agrega aquí los enlaces correspondientes.")
                doc.save(doc_path)
                print(f"Documento creado: {doc_path}")
            else:
                print(f"Ya existe: {doc_path}")
        else:
            print(f"No existe la subcarpeta: {ruta_sub}")
