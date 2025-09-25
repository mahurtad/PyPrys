import os
import time
import requests
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

# --------------------------
# CONFIGURACIÓN
# --------------------------
URLS_PRINCIPALES = [
    "https://conexion.cientifica.edu.pe/es_ES/convalidación-de-inglés",
    "https://conexion.cientifica.edu.pe/es_ES/centro-de-idiomas1",
    "https://conexion.cientifica.edu.pe/es_ES/examen-de-suficiencia-y-clasificación",
    "https://conexion.cientifica.edu.pe/es_ES/preguntas-frecuentes"
]
RUTA_WORD = r"G:\My Drive\Shared GyM\UCSUR\contenido_idiomas_cientifica.docx"
BASE_URL = "https://conexion.cientifica.edu.pe"

# --------------------------
# FUNCIONES
# --------------------------

def iniciar_driver():
    opciones = Options()
    opciones.add_argument("--headless")
    opciones.add_argument("--disable-gpu")
    opciones.add_argument("--no-sandbox")
    opciones.add_argument("--log-level=3")
    servicio = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=servicio, options=opciones)

def obtener_html_selenium(url, driver):
    driver.get(url)
    time.sleep(3)
    return driver.page_source

def limpiar_texto(texto):
    return ' '.join(texto.strip().split())

def procesar_url(url, doc, driver, visitadas):
    if url in visitadas:
        return
    visitadas.add(url)

    html = obtener_html_selenium(url, driver)
    soup = BeautifulSoup(html, "html.parser")

    doc.add_page_break()
    doc.add_heading(f"Página: {url}", level=1)

    # Título principal
    titulo = soup.title.string.strip() if soup.title else "Sin título"
    doc.add_heading(f"Título: {titulo}", level=2)

    # Encabezados
    for i in range(1, 4):
        for h in soup.find_all(f"h{i}"):
            texto = limpiar_texto(h.get_text())
            if texto:
                doc.add_heading(texto, level=i)

    # Párrafos
    for p in soup.find_all("p"):
        texto = limpiar_texto(p.get_text())
        if texto:
            doc.add_paragraph(texto)

    # Listas
    for ul in soup.find_all("ul"):
        for li in ul.find_all("li"):
            texto = limpiar_texto(li.get_text())
            if texto:
                doc.add_paragraph(f"- {texto}", style="List Bullet")

    # Enlaces internos para expandir (ej: FAQ)
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if href.startswith("/es_ES/") and not href.startswith("mailto"):
            sub_url = BASE_URL + href
            texto_link = limpiar_texto(a.get_text())
            if sub_url not in visitadas:
                doc.add_heading(f"[Subpágina] {texto_link}", level=3)
                procesar_url(sub_url, doc, driver, visitadas)

# --------------------------
# EJECUCIÓN
# --------------------------

if __name__ == "__main__":
    doc = Document()
    doc.add_heading("Contenido web: Área de Idiomas - UCSUR", 0)

    driver = iniciar_driver()
    urls_visitadas = set()

    for url in URLS_PRINCIPALES:
        procesar_url(url, doc, driver, urls_visitadas)

    driver.quit()

    os.makedirs(os.path.dirname(RUTA_WORD), exist_ok=True)
    doc.save(RUTA_WORD)

    print(f"✅ Documento guardado en: {RUTA_WORD}")
