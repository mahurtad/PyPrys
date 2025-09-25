import sys
import traceback
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = r"C:\Users\manue\Downloads\Silabo"
EXCEL_FILENAME = "Calendario-sílabo.xlsx"
TARGET_SHEET_PREFERRED = "Calendario"


def pick_sheet_with_data(xls: pd.ExcelFile) -> str:
    if TARGET_SHEET_PREFERRED in xls.sheet_names:
        return TARGET_SHEET_PREFERRED
    for sh in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sh, header=0)
        if not df.empty and df.dropna(how="all").shape[0] > 0:
            return sh
    return xls.sheet_names[0]


def dataframe_from_excel(xlsx_path: Path) -> pd.DataFrame:
    xls = pd.ExcelFile(xlsx_path, engine="openpyxl")
    sheet = pick_sheet_with_data(xls)
    df = pd.read_excel(xls, sheet_name=sheet, header=0)
    df = df.dropna(how="all").dropna(axis=1, how="all").fillna("")
    for col in df.columns:
        df[col] = df[col].astype(str)
    return df


def set_table_borders(table):
    """
    Fuerza bordes visibles en toda la tabla.
    """
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')     # grosor
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # negro
        tblBorders.append(border)
    tblPr = tbl.tblPr
    tblPr.append(tblBorders)


def add_dataframe_as_table(doc: Document, df: pd.DataFrame):
    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # encabezados
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        run = hdr_cells[j].paragraphs[0].add_run(str(col))
        run.bold = True
        run.font.size = Pt(10)

    # celdas
    for i in range(rows):
        row_cells = table.rows[i + 1].cells
        for j, col in enumerate(df.columns):
            para = row_cells[j].paragraphs[0]
            para.clear()
            run = para.add_run(str(df.iloc[i, j]))
            run.font.size = Pt(10)


def process_excel(xlsx_path: Path) -> Path:
    df = dataframe_from_excel(xlsx_path)
    doc = Document()
    doc.add_heading(xlsx_path.stem, level=1)  # título con nombre del archivo

    if df.empty:
        doc.add_paragraph("No se encontraron datos en el Excel (cuadro vacío).")
    else:
        add_dataframe_as_table(doc, df)

    out_path = xlsx_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def main():
    base = Path(BASE_DIR)
    excel_files = list(base.rglob(EXCEL_FILENAME))
    if not excel_files:
        print("No se encontraron archivos.")
        sys.exit(0)

    for xlsx in excel_files:
        try:
            print(f"Procesando {xlsx}...")
            out = process_excel(xlsx)
            print(f"✓ Generado {out}")
        except Exception as e:
            print(f"Error con {xlsx}: {e}")
            traceback.print_exc()


if __name__ == "__main__":
    main()
